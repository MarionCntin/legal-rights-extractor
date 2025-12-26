from pathlib import Path
from typing import List
from pypdf import PdfReader
from docx import Document as DocxDocument

from .normalize import Document, Page, compute_doc_id

def load_pdf(path: Path) -> Document:
    reader = PdfReader(str(path))
    pages: List[Page] = []
    for i, p in enumerate(reader.pages):
        text = p.extract_text() or ""
        pages.append(Page(page_index=i, text=text.strip()))
    return Document(
        doc_id=compute_doc_id(path),
        source_path=str(path),
        file_type="pdf",
        pages=pages,
    )

def load_docx(path: Path) -> Document:
    doc = DocxDocument(str(path))
    text = "\n".join([p.text for p in doc.paragraphs if p.text is not None]).strip()
    pages = [Page(page_index=0, text=text)]
    return Document(
        doc_id=compute_doc_id(path),
        source_path=str(path),
        file_type="docx",
        pages=pages,
    )

def load_txt(path: Path) -> Document:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    pages = [Page(page_index=0, text=text)]
    return Document(
        doc_id=compute_doc_id(path),
        source_path=str(path),
        file_type="txt",
        pages=pages,
    )

def load_any(path: Path) -> Document:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return load_pdf(path)
    if suffix == ".docx":
        return load_docx(path)
    if suffix == ".txt":
        return load_txt(path)
    raise ValueError(f"Unsupported file type: {suffix}")
